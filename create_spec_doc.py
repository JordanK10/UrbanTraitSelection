from docx import Document
from docx.shared import Inches

def create_spec_document():
    """
    Generates a .docx file containing the generalized data formatting specification.
    """
    document = Document()

    # Title
    document.add_heading('Generalized Data Formatting Specification for Spatial Analysis', level=1)

    # 1. Overview
    document.add_heading('1. Overview', level=2)
    p = document.add_paragraph()
    p.add_run('This document specifies the required data format for a generalized spatial decomposition analysis. The structure is designed to be agnostic to the specific data source (e.g., Census, proprietary data) and allows for the analysis of any spatially hierarchical system where units possess a population and a distributed characteristic (the "state variable").').italic = True
    document.add_paragraph(
        'The core idea is to track changes in a distributed state variable (like income, education level, etc.) over time across a hierarchy of nested spatial units.'
    )

    # 2. Final Data Structure
    document.add_heading('2. Final Data Structure (Pickle File)', level=2)
    document.add_paragraph(
        'The final deliverable is a single Python pickle file (.pkl) containing one dictionary.'
    )
    document.add_paragraph(
        'Top-level Keys: Integers or strings representing the time points of the data (e.g., 2014, 2023, or t0, t1).', style='List Bullet'
    )
    document.add_paragraph(
        'Nested Keys: Strings representing the spatial levels in the hierarchy, ordered from lowest to highest resolution. Example: \'level_0\', \'level_1\', \'level_2\'.', style='List Bullet'
    )
    document.add_paragraph(
        'Nested Values: A pandas DataFrame containing the data for that time point and spatial level.', style='List Bullet'
    )
    document.add_paragraph('Example Structure:').bold = True
    
    code_block = """
{
    't0': {
        'level_0': DataFrame_L0_T0,  # e.g., Block Group
        'level_1': DataFrame_L1_T0,  # e.g., Tract
        'level_2': DataFrame_L2_T0,  # e.g., County
        # ... up to the highest level
    },
    't1': {
        'level_0': DataFrame_L0_T1,
        'level_1': DataFrame_L1_T1,
        'level_2': DataFrame_L2_T1,
        # ... etc. for all time points
    }
}
"""
    document.add_paragraph(code_block, style='Normal')


    # 3. DataFrame Schema
    document.add_heading('3. DataFrame Schema & Required Columns', level=2)
    p = document.add_paragraph('Each DataFrame, regardless of its level, must adhere to the following schema. ')
    p.add_run('Consistency in column naming and data types is critical.').bold = True

    # 3.1 Identifiers
    document.add_heading('3.1. Identifiers (The "Spatial Units")', level=3)
    document.add_paragraph(
        'These columns define the units and the hierarchical structure. All identifiers must be strings or integers and must be consistent across all time points for a given unit.'
    )
    document.add_paragraph('Unit_ID (Required): A unique identifier for the spatial unit within its level. This is the primary key for each row. The DataFrame\'s index should be set to this column.', style='List Bullet')
    document.add_paragraph('Parent_ID (Required, except for the highest level): The Unit_ID of the parent unit in the level directly above. This column establishes the hierarchy. For the highest spatial level (e.g., \'level_N\'), this column can be omitted or left null.', style='List Bullet')
    
    # 3.2 Population
    document.add_heading('3.2. Population/Weighting Variable (The "Population Variables")', level=3)
    document.add_paragraph(
        'This represents the size or importance of each spatial unit.'
    )
    document.add_paragraph('Population (Required, Numeric): The total population, number of households, or other weighting factor for the unit. This value is used for calculating weighted averages and growth dynamics.', style='List Bullet')

    # 3.3 State Variable Bins
    document.add_heading('3.3. State Variable Distribution (The "State Variable Bins")', level=3)
    document.add_paragraph(
        'These columns describe the key characteristic being analyzed for each spatial unit.'
    )
    document.add_paragraph('Bin_1, Bin_2, ..., Bin_M (Required, Numeric): A set of columns representing a histogram or distribution. Each column contains the count of individuals/households within that specific bin of the state variable (e.g., number of households with income between $10k-$15k).', style='List Bullet')
    
    # Nested bullet points for the above
    p = document.add_paragraph('The number of bins, M, must be ', style='List Number 2')
    p.add_run('identical').bold = True
    p.add_run(' across all spatial units and all time points.')
    document.add_paragraph('The sum of all Bin_* columns for a given unit should represent the total number of observed entities (e.g., total households) for that unit\'s state variable.', style='List Number 2')

    # 4. Example
    document.add_heading('4. Example DataFrame (level_1 data at t0)', level=2)
    
    table_data = [
        ('Unit_ID (Index)', 'Parent_ID', 'Population', 'Bin_1', 'Bin_2', '...', 'Bin_M'),
        ('tract_A', 'county_X', '5000', '100', '150', '...', '50'),
        ('tract_B', 'county_X', '3500', '80', '120', '...', '40'),
        ('tract_C', 'county_Y', '7200', '250', '300', '...', '95'),
    ]

    table = document.add_table(rows=1, cols=len(table_data[0]))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, item in enumerate(table_data[0]):
        hdr_cells[i].text = item

    # Data rows
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data):
            row_cells[i].text = item

    # Save the document
    file_path = 'data_specification.docx'
    document.save(file_path)
    print(f"Specification document saved to {file_path}")

if __name__ == '__main__':
    create_spec_document() 